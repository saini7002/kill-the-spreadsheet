"""Kill the Quote — AI procurement analyst backend.

One end-to-end flow: draft an RFx with AI, (fake) send it to vendors, read whatever
they send back in any format, normalize every quote into one comparable basis, and
let a buyer interrogate the result in natural language toward a defensible award.

Only the plumbing is stubbed (email send). The AI loops are real: extraction and
analyst reasoning both call Google Gemini on the actual source artifacts.
"""

import os
import io
import re
import csv
import json
import base64
import mimetypes
from pathlib import Path
from datetime import datetime, timezone, date
from typing import Any

import requests
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from openpyxl import load_workbook
from pypdf import PdfReader
from docx import Document

# --------------------------------------------------------------------------- #
# Configuration
# --------------------------------------------------------------------------- #
ROOT = Path(__file__).parent
DATA = ROOT / "data"
VENDOR_DIR = DATA / "vendors"
HIST_DIR = DATA / "historical"
STATIC = ROOT / "static"

GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.5-flash-lite")
FX_USD_INR = float(os.getenv("FX_USD_INR", "95"))   # single source of truth
GST = float(os.getenv("GST_RATE", "0.18"))
TODAY = date.today().isoformat()
MAX_DELIVERY_DAYS = 30
QUOTE_VALIDITY_DAYS = 30

app = FastAPI(title="Kill the Quote — Aerchain")

with open(DATA / "requirements.json", encoding="utf-8") as f:
    REQUIREMENTS: list[dict] = json.load(f)

with open(DATA / "vendor_manifest.json", encoding="utf-8") as f:
    MANIFEST: dict = json.load(f)

VENDOR_FILES: dict[str, Path] = {
    vendor: VENDOR_DIR / meta["file"] for vendor, meta in MANIFEST.items()
}


def save_requirements() -> None:
    with open(DATA / "requirements.json", "w", encoding="utf-8") as f:
        json.dump(REQUIREMENTS, f, indent=2, ensure_ascii=False)


def next_line_id() -> str:
    nums = [int(m.group(1)) for it in REQUIREMENTS if (m := re.match(r"IT-(\d+)", str(it.get("id", ""))))]
    return f"IT-{(max(nums) + 1) if nums else 1:03d}"

# Deterministic, obviously-simulated vendor contacts for the fake email flow.
VENDOR_EMAILS: dict[str, str] = {
    "Vertex Systems": "bids@vertex-systems.example",
    "Northstar IT": "sales@northstar-it.example",
    "BluePeak Technologies": "quotes@bluepeak.example",
    "Orion Office Tech": "desk@orion-office.example",
    "TechSource India": "sales@techsource.example",
}

# The RFx the buyer builds: scope, questionnaire and commercial terms. Persisted so the
# co-pilot and manual edits survive across requests. Line items live in requirements.json.
DEFAULT_SCOPE = (
    "Company-wide employee IT hardware refresh for a 5,000-person enterprise: laptops, monitors, "
    "docks and peripherals across the line items below, sourced from five approved suppliers. "
    "INR preferred (USD accepted), GST and freight stated separately, delivery within 30 days of PO, "
    "quotes valid 30 days."
)

DEFAULT_QUESTIONNAIRE = [
    "Are you an authorized manufacturer/reseller for the quoted OEMs?",
    "Do you provide OEM-backed warranty support (state duration)?",
    "Can you deliver within 30 days of PO?",
    "Do you hold a current ISO 9001 certification (state validity)?",
    "Can you share enterprise deployment/support references?",
]

DEFAULT_COMMERCIAL_TERMS = [
    "Quote validity: 30 days from submission",
    "Currency: INR preferred; USD accepted",
    "State clearly whether GST is included or excluded",
    "State clearly whether freight/delivery is included or excluded",
    "State delivery lead time in days",
    "Quote the exact requested configuration; flag any alternative offered",
    f"Evaluation FX rate: INR {FX_USD_INR:.0f} / USD",
]


def load_rfx() -> dict:
    path = DATA / "rfx.json"
    if path.exists():
        try:
            data = json.load(open(path, encoding="utf-8"))
            return {
                "scope": data.get("scope") or DEFAULT_SCOPE,
                "questionnaire": data.get("questionnaire") or list(DEFAULT_QUESTIONNAIRE),
                "commercial_terms": data.get("commercial_terms") or list(DEFAULT_COMMERCIAL_TERMS),
            }
        except Exception:
            pass
    return {
        "scope": DEFAULT_SCOPE,
        "questionnaire": list(DEFAULT_QUESTIONNAIRE),
        "commercial_terms": list(DEFAULT_COMMERCIAL_TERMS),
    }


RFX = load_rfx()


def save_rfx() -> None:
    with open(DATA / "rfx.json", "w", encoding="utf-8") as f:
        json.dump(RFX, f, indent=2, ensure_ascii=False)



# --------------------------------------------------------------------------- #
# Historical reference data (prior-year, context only — never a current quote)
# --------------------------------------------------------------------------- #
def _wb_rows(wb) -> list[dict]:
    ws = wb["2025 Quote"] if "2025 Quote" in wb.sheetnames else wb.active
    headers = [c.value for c in ws[1]]
    rows = []
    for values in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in values):
            rows.append(dict(zip(headers, values)))
    return rows


def load_historical() -> list[dict]:
    path = HIST_DIR / "TechSource_India_2025.xlsx"
    return _wb_rows(load_workbook(path, data_only=True)) if path.exists() else []


HISTORICAL = load_historical()


def prices_from_rows(rows: list[dict]) -> dict:
    """RFx-line prices from a prior-year record, for resolving 'same as last year'."""
    out = {}
    for r in rows:
        lid = r.get("Line ID")
        price = r.get("Unit Price")
        if lid and str(lid).startswith("IT-") and isinstance(price, (int, float)):
            out[str(lid)] = {"price": float(price), "uom": (r.get("UOM") or "ea"), "currency": (r.get("Currency") or "INR")}
    return out


def questionnaire_from_rows(rows: list[dict]) -> dict:
    """Prior-year questionnaire answers for durable attributes, keyed by a lowercased label."""
    labels = ("authorized reseller", "oem warranty", "iso 9001", "enterprise references")
    out = {}
    for r in rows:
        label = str(r.get("Product / Description") or "").strip().lower()
        answer = r.get("Commercial Notes")
        if label in labels and answer:
            out[label] = str(answer)
    return out


def fy25_prices() -> dict:
    return prices_from_rows(HISTORICAL)


def fy25_questionnaire() -> dict:
    return questionnaire_from_rows(HISTORICAL)


def fy25_vendor() -> str:
    """The supplier whose prior-year record we hold (owner of the FY25 sheet)."""
    for r in HISTORICAL:
        if str(r.get("Line ID") or "").startswith("IT-") and r.get("Vendor"):
            return str(r["Vendor"])
    return ""


# --------------------------------------------------------------------------- #
# Local document parsing (text is extracted locally; PDFs/images also go to
# Gemini as native multimodal input so it can read the original layout/photo).
# --------------------------------------------------------------------------- #
def extract_local_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".xlsx":
        wb = load_workbook(path, data_only=True)
        out: list[str] = []
        for ws in wb.worksheets:
            out.append(f"[sheet: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                out.append(" | ".join("" if c is None else str(c) for c in row))
        return "\n".join(out)
    if ext == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if ext == ".docx":
        doc = Document(str(path))
        out = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                out.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(out)
    if ext in {".jpg", ".jpeg", ".png"}:
        return "[Scanned image — read the attached original image directly.]"
    return path.read_text(encoding="utf-8", errors="ignore")


# --------------------------------------------------------------------------- #
# The single deterministic normalization engine.
# Turns a supplier's raw quote into one comparable basis: pre-tax INR per unit.
# Used everywhere (comparison, CSV, analyst) so numbers can never disagree.
# --------------------------------------------------------------------------- #
def rfx_quantity(line_id: str):
    """RFx-requested quantity for a line, used to test quantity-conditional pricing."""
    for it in REQUIREMENTS:
        if it["id"] == line_id:
            return it.get("quantity")
    return None


def apply_volume_discount(price: float, q: dict) -> tuple[float, str]:
    """Apply a supplier's quantity-conditional discount only when the RFx quantity clears the threshold.

    Returns (price, note). The note records the decision either way so the buyer can see that the
    condition was evaluated, not silently ignored.
    """
    pct = q.get("volume_discount_pct")
    threshold = q.get("volume_discount_min_qty")
    if not pct or threshold is None:
        return price, ""
    try:
        pct = float(pct)
        threshold = float(threshold)
    except (TypeError, ValueError):
        return price, ""
    qty = rfx_quantity(q.get("line_id"))
    if qty is None:
        return price, ""
    if qty > threshold:
        return price * (1 - pct / 100), f"volume -{pct:g}% (qty {qty:g} over {threshold:g})"
    return price, f"volume {pct:g}% offer needs qty over {threshold:g} (have {qty:g})"


_WATT_RE = re.compile(r"(\d+)\s*W\b", re.I)


def _first_watt(text):
    m = _WATT_RE.search(text or "")
    return int(m.group(1)) if m else None


_MISBIND_STOP = {"the", "and", "for", "with", "gen", "usb", "c", "w", "pd", "ips", "fhd", "4k", "it", "ea", "inch", "adjustable", "height"}


def _desc_tokens(s: str) -> set:
    return {t for t in re.findall(r"[a-z0-9]+", (s or "").lower()) if len(t) > 1 and t not in _MISBIND_STOP}


def flag_misbound(q: dict) -> None:
    """Catch a scan row-shift: if a line's extracted content clearly matches a DIFFERENT RFx line's
    description than the one it is bound to, the row was mis-read (common with angled photos). Discard it
    as MISSING rather than trust the shifted value. General rule, not tied to any one line id."""
    if q.get("status") != "QUOTED":
        return
    ev = _desc_tokens(q.get("evidence"))
    if not ev:
        return
    scores = {it["id"]: len(ev & _desc_tokens(it["description"])) for it in REQUIREMENTS}
    own = scores.get(q.get("line_id"), 0)
    best_id = max(scores, key=scores.get)
    if best_id != q.get("line_id") and scores[best_id] >= own + 3:
        q["status"] = "MISSING"
        q["unit_price"] = None
        q["spec_mismatch"] = False
        q["spec_mismatch_note"] = None
        q["evidence"] = f"no reliable row for this line (the extracted row matched {best_id}, likely a scan mis-read)"


def flag_spec_substitute(q: dict) -> None:
    """Deterministic backstop for alternate configurations.

    If the RFx line and the supplier's own transcribed spec both cite a wattage and they differ,
    mark the line a substitute even when the extractor did not. General rule (compares like specs
    on the same line), not tied to any one line id.
    """
    if q.get("status") != "QUOTED" or q.get("spec_mismatch"):
        return
    rfx = next((it for it in REQUIREMENTS if it["id"] == q.get("line_id")), None)
    if not rfx:
        return
    rw = _first_watt(rfx.get("description"))
    vw = _first_watt(q.get("evidence"))
    if rw and vw and rw != vw:
        q["spec_mismatch"] = True
        q["spec_mismatch_note"] = q.get("spec_mismatch_note") or f"supplier quotes {vw}W; RFx line requires {rw}W"


_UNIT_TOKENS = re.compile(r"/\s*100|per\s*100|100\s*pc|/\s*box|per\s*box|/\s*pack|per\s*pack|per\s*unit|each|/\s*ea|\bea\b|\bpcs?\b|piece", re.I)


def flag_unit_inheritance(vendor: dict) -> None:
    """When a vendor prices some lines per a pack unit (per-100 / box), a QUOTED line whose evidence
    states NO unit of its own has an ambiguous unit (it could be per pack or per unit). Flag it, and set
    the assumption type deterministically so the UI offers the right action."""
    quotes = [q for q in vendor.get("quotes", []) if q.get("status") == "QUOTED"]
    has_pack = any(any(k in (q.get("quoted_uom") or "").lower() for k in ("100", "box", "pack", "carton", "dozen")) for q in quotes)
    if not has_pack:
        return
    for q in quotes:
        if q.get("provenance") or _UNIT_TOKENS.search(q.get("evidence") or ""):
            continue
        q["assumption"] = "unit not stated by the supplier; a nearby line was priced per pack, so this could be per pack rather than per unit"
        q["assumption_type"] = "unit"


def flag_discount_scope(vendor: dict) -> None:
    """A discount clause shared across several lines is only ambiguous for a line it did NOT directly
    follow. We flag a discount-bearing line only when the discount was actually APPLIED to it (materially
    changed its price) and it is part of a multi-line clause; an unapplied or single-line discount is
    never flagged. Assigned deterministically (overrides any misplaced model assumption)."""
    disc = [q for q in vendor.get("quotes", []) if q.get("volume_discount_pct") and q.get("status") == "QUOTED"]
    if len(disc) < 2:
        return
    for q in disc:
        adj = (q.get("normalization") or {}).get("adjustments") or []
        applied = any(str(a).startswith("volume -") for a in adj)
        if applied and not q.get("provenance"):
            pct = q["volume_discount_pct"]
            q["assumption"] = (f"a {pct:g}% discount clause listed near several lines was applied to this line; "
                               "confirm it covers this line, as it may apply only to the line it directly follows")
            q["assumption_type"] = "discount"


def flag_scale_outliers(vendors: list[dict]) -> None:
    """Cross-vendor sanity check: a comparable line priced far above the next quote on the same line is
    almost always a unit/scale error (e.g. a per-100 or per-box price read as per-unit). Flag it for
    review rather than trusting it. General rule, keyed on nothing but the numbers."""
    for it in REQUIREMENTS:
        lid = it["id"]
        entries = []
        for v in vendors:
            for q in v.get("quotes", []):
                n = (q.get("normalization") or {}).get("normalized_unit_price_inr")
                if q.get("line_id") == lid and q.get("status") == "QUOTED" and (q.get("normalization") or {}).get("comparable") and n is not None:
                    entries.append((n, q))
        if len(entries) < 2:
            continue
        lo = min(n for n, _ in entries)
        if not lo:
            continue
        for n, q in entries:
            if n > lo * 10 and not q.get("provenance"):
                note = f"price of {n:.0f} is far above the next quote ({lo:.0f}) on this line; likely a unit or scale mismatch"
                q["assumption"] = note + ((" | " + q["assumption"]) if q.get("assumption") else "")
                q["assumption_type"] = q.get("assumption_type") or "scale"


def normalize_quote(q: dict) -> dict:
    adjustments: list[str] = []
    result = {
        "normalized_unit_price_inr": None,
        "adjustments": adjustments,
        "comparable": False,
        "freight_extra": q.get("freight_included") is False,
    }
    if q.get("status") != "QUOTED" or q.get("unit_price") is None:
        return result

    price = float(q["unit_price"])
    currency = (q.get("currency") or "INR").upper()
    uom = (q.get("quoted_uom") or "").lower()

    if currency == "USD":
        price *= FX_USD_INR
        adjustments.append(f"USD→INR @ {FX_USD_INR:.0f}")

    if any(k in uom for k in ("100 pc", "100pc", "100 piece", "per 100", "/100", "100 nos")):
        price /= 100
        adjustments.append("per-100 → per-unit")
    elif any(k in uom for k in ("box", "pack", "carton", "dozen")):
        pack = q.get("pack_quantity")
        try:
            pack = float(pack) if pack is not None else 0.0
        except (TypeError, ValueError):
            pack = 0.0
        if pack > 0:
            price /= pack
            adjustments.append(f"per pack of {pack:g} → per-unit")
        else:
            # Priced per box/pack but the pack size is unknown: cannot reduce to a per-unit basis.
            adjustments.append("priced per box but pack size not stated: cannot compare")
            return result

    if q.get("tax_included") is True:
        price /= (1 + GST)
        adjustments.append(f"removed {GST * 100:.0f}% GST")

    price, vol_note = apply_volume_discount(price, q)
    if vol_note:
        adjustments.append(vol_note)

    result["normalized_unit_price_inr"] = round(price, 2)
    # An alternate configuration keeps its price visible but is not comparable to the RFx line
    # until the buyer accepts equivalence.
    if q.get("spec_mismatch") is True:
        adjustments.append(f"substitute: {q.get('spec_mismatch_note') or 'alternate configuration'}")
        result["comparable"] = False
    else:
        result["comparable"] = True
    return result


def warranty_months(text) -> int | None:
    """Best-effort parse of a free-text warranty into months (e.g. '3 years' -> 36, '12 months' -> 12)."""
    if not text:
        return None
    m = re.search(r"(\d+)\s*year", str(text), re.I)
    if m:
        return int(m.group(1)) * 12
    m = re.search(r"(\d+)\s*month", str(text), re.I)
    if m:
        return int(m.group(1))
    return None


def derive_qual_status(vendor: dict) -> None:
    """Overall qualification from per-item results: any NOT_MET fails; any UNCLEAR is UNKNOWN; else PASS."""
    results = [i.get("result") for i in (vendor.get("qualification") or {}).get("items") or []]
    if "NOT_MET" in results:
        vendor["qualification"]["status"] = "FAIL"
    elif "UNCLEAR" in results:
        vendor["qualification"]["status"] = "UNKNOWN"
    elif results:
        vendor["qualification"]["status"] = "PASS"


def derive_commercial_terms(vendor: dict) -> None:
    """Backfill vendor-level commercial terms from the quotes and derive buyer risk flags deterministically.

    The model records facts only; the SYSTEM decides what is a risk. A term is flagged only when it
    genuinely disadvantages the buyer: a quote validity shorter than the RFx window (too little time to
    raise a PO) or restricted stock/availability. Payment days, freight/tax and warranty are neutral
    comparison terms and never appear as risks.
    """
    terms = vendor.get("commercial_terms")
    if not isinstance(terms, dict):
        terms = {}
        vendor["commercial_terms"] = terms
    quotes = vendor.get("quotes", [])

    def _majority_bool(field):
        vals = [q.get(field) for q in quotes if q.get(field) is not None]
        if not vals:
            return None
        return sum(1 for v in vals if v) >= (len(vals) / 2)

    if terms.get("freight_included") is None:
        terms["freight_included"] = _majority_bool("freight_included")
    if terms.get("taxes_included") is None:
        terms["taxes_included"] = _majority_bool("tax_included")
    if terms.get("delivery_days") is None:
        leads = [q.get("lead_time_days") for q in quotes if q.get("lead_time_days") is not None]
        terms["delivery_days"] = max(leads) if leads else None

    risks = []
    val = terms.get("quote_validity_days")
    if isinstance(val, (int, float)) and val < QUOTE_VALIDITY_DAYS:
        risks.append(f"quote valid {int(val)} days only, may be too short to raise a PO")
    avail = terms.get("availability")
    if isinstance(avail, str) and avail.strip():
        risks.append(f"{avail.strip()}, confirm stock before relying on this supplier")
    vendor["commercial_risks"] = risks


def enrich_vendor(vendor: dict) -> dict:
    """Attach normalized numbers to every quote, and derive qualification status from its items.

    Overall status is computed deterministically so every questionnaire item counts equally:
    any NOT_MET fails the supplier; an unresolved item leaves it UNKNOWN; otherwise PASS.
    """
    for q in vendor.get("quotes", []):
        flag_misbound(q)
        flag_spec_substitute(q)
        q["normalization"] = normalize_quote(q)
        # Discard the model's free-text assumption; assumptions are assigned deterministically below.
        q["assumption"] = ""
        q["assumption_type"] = None
    flag_unit_inheritance(vendor)
    flag_discount_scope(vendor)
    derive_qual_status(vendor)
    derive_commercial_terms(vendor)
    return vendor


def compute_award_context(extracted: dict) -> dict:
    """Deterministically compute award totals so the analyst cites real numbers, not guesses.

    Only quoted, comparable lines from qualified (PASS) suppliers are award-eligible.
    """
    vendors = extracted.get("vendors", [])
    qual = {v.get("vendor"): v.get("qualification", {}).get("status") for v in vendors}
    qty = {it["id"]: it.get("quantity", 0) for it in REQUIREMENTS}
    bundle = {v.get("vendor"): (v.get("bundle_discount_pct") or 0) for v in vendors}

    price: dict[tuple[str, str], float] = {}
    quote_by: dict[tuple[str, str], dict] = {}
    for v in vendors:
        for q in v.get("quotes", []):
            key = (v.get("vendor"), q.get("line_id"))
            quote_by[key] = q
            norm = q.get("normalization") or {}
            unit = norm.get("normalized_unit_price_inr")
            if q.get("status") == "QUOTED" and norm.get("comparable") and unit is not None:
                price[key] = unit

    # Split award: cheapest qualified comparable price on each line.
    split_total = 0.0
    split_lines = []
    split_by_vendor: dict[str, dict] = {}
    unavailable = []
    for it in REQUIREMENTS:
        lid = it["id"]
        options = [(price[(vn, lid)], vn) for vn in qual if qual[vn] == "PASS" and (vn, lid) in price]
        if not options:
            unavailable.append(lid)
            continue
        unit, vn = min(options)
        line_total = round(unit * qty[lid], 2)
        split_total += line_total
        split_lines.append({"line_id": lid, "vendor": vn, "unit_price_inr": unit, "quantity": qty[lid], "line_total_inr": line_total})
        b = split_by_vendor.setdefault(vn, {"lines": 0, "subtotal_inr": 0.0, "carried_lines": 0, "buyer_edited_lines": 0, "fresh_quote_lines": 0})
        b["lines"] += 1
        b["subtotal_inr"] = round(b["subtotal_inr"] + line_total, 2)
        # Firmness of THIS vendor's split-winning lines (the in-context carried count, not its total).
        prov = (quote_by.get((vn, lid)) or {}).get("provenance")
        if prov == "carried":
            b["carried_lines"] += 1
        elif prov == "buyer-edited":
            b["buyer_edited_lines"] += 1
        else:
            b["fresh_quote_lines"] += 1

    # Single-supplier award: per qualified vendor, only the lines it can comparably serve.
    single = []
    for v in vendors:
        vn = v.get("vendor")
        if qual.get(vn) != "PASS":
            continue
        subtotal = 0.0
        covered = []
        missing = []
        for it in REQUIREMENTS:
            lid = it["id"]
            if (vn, lid) in price:
                subtotal += price[(vn, lid)] * qty[lid]
                covered.append(lid)
            else:
                missing.append(lid)
        entry = {
            "vendor": vn,
            "lines_covered": len(covered),
            "lines_missing": len(missing),
            "missing_line_ids": missing,
            "subtotal_inr": round(subtotal, 2),
            "note": "subtotal covers only the lines this supplier can comparably serve; not a full-basket total",
        }
        # A full-RFx bundle discount only bites when the entire order goes to this one supplier.
        bpct = bundle.get(vn) or 0
        if bpct:
            disc = round(subtotal * bpct / 100, 2)
            entry["bundle_discount_pct"] = bpct
            entry["bundle_discount_inr"] = disc
            entry["subtotal_after_discount_inr"] = round(subtotal - disc, 2)
            note = f"; includes {bpct:g}% full-RFx bundle discount, applied only in this single-supplier scenario"
            if missing:
                note += f" (assumes the {len(missing)} uncovered line(s) are also resolved with this supplier)"
            entry["note"] += note
        single.append(entry)

    # Explicit split-vs-single comparison so the analyst never has to reason out which is cheaper.
    # Only suppliers that can serve every line are eligible for a full-basket single-supplier total.
    full_basket = [
        (s.get("subtotal_after_discount_inr", s["subtotal_inr"]), s["vendor"], "subtotal_after_discount_inr" in s)
        for s in single if s["lines_missing"] == 0
    ]
    comparison = {"split_total_inr": round(split_total, 2)}
    if full_basket:
        best_total, best_vendor, has_bundle = min(full_basket)
        comparison["cheapest_full_basket_single_supplier"] = {
            "vendor": best_vendor,
            "total_inr": best_total,
            "includes_bundle_discount": has_bundle,
        }
        if best_total < split_total:
            comparison["lowest_cost_strategy"] = f"single-supplier: {best_vendor}"
            comparison["single_supplier_is_lower_than_split_by_inr"] = round(split_total - best_total, 2)
        elif best_total > split_total:
            comparison["lowest_cost_strategy"] = "split award"
            comparison["split_is_lower_than_cheapest_single_by_inr"] = round(best_total - split_total, 2)
        else:
            comparison["lowest_cost_strategy"] = "split award and cheapest single-supplier are equal"
        comparison["note"] = ("Compare only these exact totals to state which strategy is cheaper; "
                              "do not infer the direction yourself.")

    # Warranty-filtered award: cheapest way to buy if the buyer insists on a minimum warranty tier,
    # and the premium over the overall cheapest option. Computed here so the analyst never estimates it.
    WARRANTY_TIER_MONTHS = 36
    wmonths = {v.get("vendor"): warranty_months((v.get("commercial_terms") or {}).get("warranty")) for v in vendors}
    tier_vendors = [vn for vn in qual
                    if qual[vn] == "PASS" and wmonths.get(vn) is not None and wmonths[vn] >= WARRANTY_TIER_MONTHS]
    warranty_alternatives = {"threshold_months": WARRANTY_TIER_MONTHS, "vendors_meeting_threshold": tier_vendors}
    if tier_vendors:
        # Cheapest split restricted to tier vendors.
        w_split_total, w_split_missing, w_split_by_vendor = 0.0, [], {}
        for it in REQUIREMENTS:
            lid = it["id"]
            opts = [(price[(vn, lid)], vn) for vn in tier_vendors if (vn, lid) in price]
            if not opts:
                w_split_missing.append(lid)
                continue
            u, vn = min(opts)
            w_split_total += u * qty[lid]
            w_split_by_vendor[vn] = w_split_by_vendor.get(vn, 0) + 1
        # Cheapest single tier vendor covering every line (bundle applied if offered).
        w_singles = []
        for vn in tier_vendors:
            if all((vn, it["id"]) in price for it in REQUIREMENTS):
                sub = sum(price[(vn, it["id"])] * qty[it["id"]] for it in REQUIREMENTS)
                bpct = bundle.get(vn) or 0
                w_singles.append((round(sub * (1 - bpct / 100), 2), vn))
        baseline = comparison.get("cheapest_full_basket_single_supplier", {})
        baseline_total = min(baseline.get("total_inr", split_total), round(split_total, 2))
        if not w_split_missing:
            warranty_alternatives["cheapest_split_inr"] = round(w_split_total, 2)
            warranty_alternatives["cheapest_split_by_vendor_lines"] = w_split_by_vendor
            warranty_alternatives["premium_split_over_cheapest_inr"] = round(w_split_total - baseline_total, 2)
            warranty_alternatives["premium_split_over_cheapest_pct"] = round((w_split_total - baseline_total) / baseline_total * 100, 1)
        if w_singles:
            wt, wv = min(w_singles)
            warranty_alternatives["cheapest_single_full_basket"] = {"vendor": wv, "total_inr": wt}
            warranty_alternatives["premium_single_over_cheapest_inr"] = round(wt - baseline_total, 2)
            warranty_alternatives["premium_single_over_cheapest_pct"] = round((wt - baseline_total) / baseline_total * 100, 1)
        warranty_alternatives["baseline_cheapest"] = {
            "vendor": baseline.get("vendor"), "total_inr": baseline_total,
            "warranty_months": wmonths.get(baseline.get("vendor")),
        }
        warranty_alternatives["note"] = ("Premiums are versus the overall cheapest option; use these exact "
                                         "figures for any 'how much more for a longer warranty' question.")

    # Per-vendor freight/tax inclusion echoed so the analyst never misstates who charges extra.
    commercial_flags = [
        {"vendor": v.get("vendor"),
         "freight_included": (v.get("commercial_terms") or {}).get("freight_included"),
         "taxes_included": (v.get("commercial_terms") or {}).get("taxes_included")}
        for v in vendors
    ]

    # Per-line spec-vs-substitute analysis so the analyst compares against the cheapest EXACT-spec
    # qualified quote instead of scanning quotes itself (which can drop the cheapest exact quote).
    spec_watch = []
    for it in REQUIREMENTS:
        lid = it["id"]
        exact, substitutes = [], []
        for v in vendors:
            vn = v.get("vendor")
            for q in v.get("quotes", []):
                if q.get("line_id") != lid:
                    continue
                norm = q.get("normalization") or {}
                unit = norm.get("normalized_unit_price_inr")
                if q.get("substitute_accepted") or q.get("spec_mismatch"):
                    substitutes.append({
                        "vendor": vn, "unit_price_inr": unit, "qualified": qual.get(vn) == "PASS",
                        "accepted": bool(q.get("substitute_accepted")),
                        "note": q.get("spec_mismatch_note") or "",
                    })
                elif q.get("status") == "QUOTED" and norm.get("comparable") and unit is not None and qual.get(vn) == "PASS":
                    exact.append((unit, vn))
        if substitutes:
            cheapest_exact = None
            if exact:
                u, vn = min(exact)
                cheapest_exact = {"vendor": vn, "unit_price_inr": u}
            spec_watch.append({"line_id": lid, "cheapest_exact_qualified": cheapest_exact, "substitutes": substitutes})

    # Deterministic per-vendor tally of buyer-resolved lines, so the analyst cites exact counts
    # (e.g. how many lines are carried) instead of tallying the dataset itself and miscounting.
    resolution_summary = []
    for v in vendors:
        quotes = v.get("quotes", [])
        carried = [q.get("line_id") for q in quotes if q.get("provenance") == "carried"]
        edited = [q.get("line_id") for q in quotes if q.get("provenance") == "buyer-edited"]
        substitutes = [q.get("line_id") for q in quotes if q.get("substitute_accepted")]
        unavailable_lines = [q.get("line_id") for q in quotes if q.get("provenance") == "buyer-unavailable"]
        carried_qual = [it.get("question") for it in (v.get("qualification") or {}).get("items", [])
                        if it.get("provenance") == "carried"]
        confirmed_qual = [it.get("question") for it in (v.get("qualification") or {}).get("items", [])
                          if it.get("provenance") == "buyer-confirmed"]
        if any([carried, edited, substitutes, unavailable_lines, carried_qual, confirmed_qual]):
            resolution_summary.append({
                "vendor": v.get("vendor"),
                "carried_lines": len(carried),
                "carried_line_ids": carried,
                "buyer_edited_lines": len(edited),
                "buyer_edited_line_ids": edited,
                "accepted_substitute_lines": len(substitutes),
                "accepted_substitute_line_ids": substitutes,
                "buyer_marked_unavailable_lines": len(unavailable_lines),
                "carried_qualification_items": carried_qual,
                "buyer_confirmed_qualification_items": confirmed_qual,
            })

    return {
        "fx_rate_usd_inr": FX_USD_INR,
        "gst_rate": GST,
        "basis": "pre-tax INR per unit; only quoted, comparable lines from qualified (PASS) suppliers are eligible",
        "total_rfx_lines": len(REQUIREMENTS),
        "split_award": {
            "strategy": "cheapest qualified comparable price per line",
            "total_inr": round(split_total, 2),
            "lines_awarded": len(split_lines),
            "lines_unavailable": unavailable,
            "per_vendor": split_by_vendor,
            "lines": split_lines,
        },
        "single_supplier_awards": single,
        "comparison": comparison,
        "warranty_alternatives": warranty_alternatives,
        "commercial_flags": commercial_flags,
        "spec_watch": spec_watch,
        "resolution_summary": resolution_summary,
    }



# --------------------------------------------------------------------------- #
# Gemini client
# --------------------------------------------------------------------------- #
def get_gemini_key() -> str:
    key = os.getenv("GEMINI_API_KEY")
    if key:
        return key.strip()
    for candidate in (ROOT / "Gemini API Key.txt", ROOT.parent / "Gemini API Key.txt"):
        if candidate.exists():
            first = candidate.read_text(encoding="utf-8").strip().splitlines()
            if first:
                return first[0].strip()
    raise HTTPException(400, "GEMINI_API_KEY is not configured.")


def gemini(parts: list[dict], schema: dict | None = None) -> str:
    payload: dict[str, Any] = {
        "contents": [{"role": "user", "parts": parts}],
        "generationConfig": {"temperature": 0.1},
    }
    if schema:
        payload["generationConfig"]["responseMimeType"] = "application/json"
        payload["generationConfig"]["responseSchema"] = schema

    resp = requests.post(
        f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent",
        params={"key": get_gemini_key()},
        headers={"Content-Type": "application/json"},
        json=payload,
        timeout=180,
    )
    if resp.status_code >= 400:
        raise HTTPException(resp.status_code, resp.text[:2000])
    data = resp.json()
    try:
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError):
        raise HTTPException(502, f"Unexpected Gemini response: {json.dumps(data)[:2000]}")


def gemini_json(parts: list[dict], schema: dict, what: str) -> Any:
    text = gemini(parts, schema)
    try:
        return json.loads(text)
    except json.JSONDecodeError as e:
        raise HTTPException(502, f"Invalid {what} from model: {e}")


def text_part(text: str) -> dict:
    return {"text": text}


IMAGE_EXTS = {".jpg", ".jpeg", ".png"}


def deskew_image_bytes(path: Path) -> tuple[bytes, str, float]:
    """Auto-straighten an angled scan so text rows align before extraction.

    Detects the rotation via horizontal projection-profile sharpness and rotates a copy.
    The source file is never modified. Returns (bytes, mime, applied_angle_degrees).
    """
    try:
        from PIL import Image
        import numpy as np

        img = Image.open(path).convert("RGB")
        mono = Image.fromarray(((np.asarray(img.convert("L"), dtype=np.float32) < 128) * 255).astype("uint8"))
        if mono.width > 700:
            mono = mono.resize((700, max(1, round(mono.height * 700 / mono.width))))

        def sharpness(angle: float) -> float:
            rot = np.asarray(mono.rotate(angle, resample=Image.BILINEAR, fillcolor=0), dtype=np.float32)
            return float(np.var(rot.sum(axis=1)))

        best = max((a / 4.0 for a in range(-24, 25)), key=sharpness)
        if abs(best) < 0.4:
            return path.read_bytes(), mimetypes.guess_type(path.name)[0] or "image/jpeg", 0.0
        out = img.rotate(best, resample=Image.BICUBIC, fillcolor=(255, 255, 255))
        buf = io.BytesIO()
        out.save(buf, format="JPEG", quality=92)
        return buf.getvalue(), "image/jpeg", round(best, 2)
    except Exception:
        return path.read_bytes(), mimetypes.guess_type(path.name)[0] or "application/octet-stream", 0.0


def file_part(path: Path) -> dict:
    if path.suffix.lower() in IMAGE_EXTS:
        data, mime, _ = deskew_image_bytes(path)
    else:
        mime = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        data = path.read_bytes()
    return {"inline_data": {"mime_type": mime, "data": base64.b64encode(data).decode("ascii")}}


def build_parts(system: str, user: str, attach: Path | None = None) -> list[dict]:
    parts = [text_part(f"{system}\n\n{user}")]
    if attach and attach.suffix.lower() in {".pdf"} | IMAGE_EXTS:
        parts.append(file_part(attach))
    return parts


# --------------------------------------------------------------------------- #
# Schemas
# --------------------------------------------------------------------------- #
def copilot_schema() -> dict:
    return {
        "type": "object",
        "properties": {
            "action_summary": {"type": "string"},
            "target": {"type": "string", "enum": ["scope", "line_items", "questionnaire", "commercial_terms", "none"]},
            "scope": {"type": "string", "nullable": True},
            "questionnaire": {"type": "array", "items": {"type": "string"}, "nullable": True},
            "commercial_terms": {"type": "array", "items": {"type": "string"}, "nullable": True},
            "line_item": {
                "type": "object",
                "nullable": True,
                "properties": {
                    "category": {"type": "string"},
                    "description": {"type": "string"},
                    "uom": {"type": "string"},
                    "quantity": {"type": "number"},
                },
                "required": ["category", "description", "uom", "quantity"],
            },
        },
        "required": ["action_summary", "target", "scope", "questionnaire", "commercial_terms", "line_item"],
    }


def extraction_schema() -> dict:
    quote = {
        "type": "object",
        "properties": {
            "line_id": {"type": "string"},
            "description": {"type": "string"},
            "quantity": {"type": "number", "nullable": True},
            "unit": {"type": "string", "nullable": True},
            "quoted_uom": {"type": "string", "nullable": True},
            "pack_quantity": {"type": "number", "nullable": True},
            "volume_discount_pct": {"type": "number", "nullable": True},
            "volume_discount_min_qty": {"type": "number", "nullable": True},
            "spec_mismatch": {"type": "boolean", "nullable": True},
            "spec_mismatch_note": {"type": "string", "nullable": True},
            "assumption": {"type": "string", "nullable": True},
            "currency": {"type": "string", "nullable": True},
            "unit_price": {"type": "number", "nullable": True},
            "tax_included": {"type": "boolean", "nullable": True},
            "freight_included": {"type": "boolean", "nullable": True},
            "lead_time_days": {"type": "number", "nullable": True},
            "status": {"type": "string", "enum": ["QUOTED", "MISSING", "AMBIGUOUS"]},
            "evidence": {"type": "string"},
            "confidence": {"type": "number"},
        },
        "required": [
            "line_id", "description", "quantity", "unit", "quoted_uom", "pack_quantity",
            "volume_discount_pct", "volume_discount_min_qty",
            "spec_mismatch", "spec_mismatch_note", "assumption",
            "currency", "unit_price", "tax_included", "freight_included", "lead_time_days",
            "status", "evidence", "confidence",
        ],
    }
    return {
        "type": "object",
        "properties": {
            "vendor": {"type": "string"},
            "qualification": {
                "type": "object",
                "properties": {
                    "status": {"type": "string", "enum": ["PASS", "FAIL", "UNKNOWN"]},
                    "evidence": {"type": "string"},
                    "items": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "question": {"type": "string"},
                                "result": {"type": "string", "enum": ["MET", "NOT_MET", "UNCLEAR"]},
                                "evidence": {"type": "string"},
                            },
                            "required": ["question", "result", "evidence"],
                        },
                    },
                },
                "required": ["status", "evidence", "items"],
            },
            "commercial_terms": {
                "type": "object",
                "properties": {
                    "warranty": {"type": "string", "nullable": True},
                    "delivery_days": {"type": "number", "nullable": True},
                    "payment_days": {"type": "number", "nullable": True},
                    "freight_included": {"type": "boolean", "nullable": True},
                    "taxes_included": {"type": "boolean", "nullable": True},
                    "quote_validity_days": {"type": "number", "nullable": True},
                    "availability": {"type": "string", "nullable": True},
                },
                "required": [
                    "warranty", "delivery_days", "payment_days", "freight_included",
                    "taxes_included", "quote_validity_days", "availability",
                ],
            },
            "bundle_discount_pct": {"type": "number", "nullable": True},
            "bundle_discount_condition": {"type": "string", "nullable": True},
            "quotes": {"type": "array", "items": quote},
            "historical": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "line_id": {"type": "string", "nullable": True},
                        "description": {"type": "string"},
                        "price": {"type": "number", "nullable": True},
                        "currency": {"type": "string", "nullable": True},
                        "note": {"type": "string"},
                    },
                    "required": ["line_id", "description", "price", "currency", "note"],
                },
            },
        },
        "required": ["vendor", "qualification", "commercial_terms", "bundle_discount_pct", "bundle_discount_condition", "quotes", "historical"],
    }


def line_item_schema() -> dict:
    return {
        "type": "object",
        "properties": {
            "category": {"type": "string"},
            "description": {"type": "string"},
            "uom": {"type": "string"},
            "quantity": {"type": "number"},
        },
        "required": ["category", "description", "uom", "quantity"],
    }


def analyst_schema() -> dict:
    return {
        "type": "object",
        "properties": {
            "answer": {"type": "string"},
            "table": {
                "type": "object",
                "nullable": True,
                "properties": {
                    "columns": {"type": "array", "items": {"type": "string"}},
                    "rows": {"type": "array", "items": {"type": "array", "items": {"type": "string"}}},
                },
                "required": ["columns", "rows"],
            },
            "chart": {
                "type": "object",
                "nullable": True,
                "properties": {
                    "type": {"type": "string", "enum": ["bar", "none"]},
                    "title": {"type": "string"},
                    "labels": {"type": "array", "items": {"type": "string"}},
                    "values": {"type": "array", "items": {"type": "number"}},
                    "unit": {"type": "string"},
                },
                "required": ["type", "title", "labels", "values", "unit"],
            },
            "warnings": {"type": "array", "items": {"type": "string"}},
            "recommended_action": {"type": "string"},
        },
        "required": ["answer", "table", "chart", "warnings", "recommended_action"],
    }


# --------------------------------------------------------------------------- #
# Prompts
# --------------------------------------------------------------------------- #
COPILOT_SYSTEM = (
    "You are an enterprise procurement RFx co-pilot. The buyer builds ONE RFx by talking to you, "
    "and each message addresses ONE part of it: scope, line_items, questionnaire, or commercial_terms. "
    "Read the buyer's message and the CURRENT RFx STATE, decide which part they mean, and return the "
    "UPDATED content for that part only, leaving the other fields null.\n"
    "- scope: return the full updated scope paragraph.\n"
    "- questionnaire: return the FULL updated list of questions (apply the buyer's add/edit/remove to the current list).\n"
    "- commercial_terms: return the FULL updated list of commercial terms.\n"
    "- line_items: return a single line_item to ADD (category, description, uom usually 'ea', quantity). "
    "The system assigns the ID; do not include one.\n"
    "- If the message does not clearly map to a part, or only asks a question, set target='none' and explain in action_summary.\n"
    "Keep action_summary to one short sentence confirming exactly what changed. Do not invent unrelated "
    "content or unnecessary procurement complexity."
)

EXTRACT_SYSTEM = """You are a meticulous enterprise procurement document extraction engine.

Extract and structure a supplier's response to the supplied RFx.

CORE PRINCIPLES
- Extract only what the supplier's current response supports. Never fabricate or silently fill gaps.
- Preserve the supplier's actual wording in the evidence field.
- Map supplier products to the canonical RFx line IDs. Do not invent line IDs.
- Return ALL RFx lines, including lines the supplier did not quote (status MISSING).

MULTIPLE FILES PER SUPPLIER
- A supplier's response may arrive as SEVERAL sources: the email body plus one or more attachments,
  and possibly a buyer-uploaded file. Read them ALL together and decide each source's role from its
  content: current quote, questionnaire answers, an ignorable cover note, or a PRIOR-YEAR / historical
  price list.
- The current quote may be in the email body itself, in an attachment, or split across several files.
  Combine them into one coherent current response.
- If a source is clearly a PRIOR-YEAR or historical price list (titled/dated a previous year, or the body
  refers to it as last year's pricing), DO NOT treat its numbers as current quotes. Put those rows in the
  `historical` array (line_id if identifiable, description, price, currency, and a short note such as
  "2025 price list"). Historical numbers must NEVER become current quotes.
- If the current response says "same as last year" for a line with no current number, that line is
  AMBIGUOUS; you may reference the matching historical row as context, but it still needs confirmation.
- If no historical source is present, return an empty historical array.

PRICE AND UNIT
- Capture the numeric price exactly as stated. Do not convert, divide, or multiply during extraction.
- Numbers may use Indian lakh grouping: "2,40,000" means 240000 and "3,50,000" means 350000. Parse the
  FULL integer and never drop digits or pre-divide.
- Capture the supplier's quoted unit of measure (quoted_uom) separately from the RFx unit.
- Capture pack_quantity only when the supplier explicitly states one (e.g. box of 10, per 100 pieces).
- If a price is quoted per box / pack / carton / dozen and the supplier does NOT state how many pieces
  the pack holds, still record the price and quoted_uom with status QUOTED and leave pack_quantity null.
  Never guess a pack size; the system will mark the line not-comparable and surface the ambiguity.
- Preserve the original currency. Keep each price bound to the exact line it belongs to.

CONDITIONAL / VOLUME DISCOUNTS
- A supplier may offer a discount that depends on order quantity for specific lines. It may appear
  inline beside a rate (e.g. a parenthetical like "5% off above 250 units") or as a separate sentence
  or footnote that names line IDs (e.g. "lines A and B: 8% off for orders over 100 units").
- For EVERY line the offer names, record volume_discount_pct (the percentage) and
  volume_discount_min_qty (the quantity threshold) on that line's quote, even when the offer text sits
  away from the price line. Keep unit_price as the full list price.
- Do NOT apply the discount yourself; the system applies it deterministically against the RFx quantity.
  Leave both fields null on lines the offer does not cover.

LINE STATUS
- QUOTED: a current numeric price can be tied to the RFx line (including an alternate configuration -
  keep it QUOTED and flag it under SPECS AND SUBSTITUTES rather than hiding it).
- MISSING: the line is absent from the current response.
- AMBIGUOUS: information may relate to the line but is not a reliable current quote
  (e.g. "same as last year" with no current number).
- A blanket catch-all such as "rest same as last year" ADDRESSES every remaining RFx line: mark each of
  those lines AMBIGUOUS (it references prior-year pricing and needs confirmation), NOT MISSING. Reserve
  MISSING only for lines the supplier does not address at all.
- A current explicit price is a current quote even if the supplier also cites history.

SPECS AND SUBSTITUTES
- Transcribe the supplier's OWN product description and specification into evidence exactly as stated.
  NEVER replace it with the RFx line's wording. If the supplier quotes a different variant, record what
  the supplier actually wrote, not what the RFx line asks for.
- Compare the key specification of the quoted item against the RFx line (for example power/PD wattage,
  screen size, storage capacity, model generation). If a key spec DIFFERS from what the RFx line
  requires, set spec_mismatch=true and spec_mismatch_note stating both sides plainly
  (e.g. "supplier quotes a 30W adapter; RFx line requires a 60W adapter"). Keep status QUOTED and keep the price;
  the system treats a mismatch as a SUBSTITUTE, not comparable, until the buyer accepts equivalence.
- If the key specs match, set spec_mismatch=false.

INFERENCE AND ASSUMPTIONS
- Reserve `assumption` for MATERIAL uncertainties only: an inherited or unclear unit of measure, the
  scope of a discount, or an alternate-spec equivalence. Do NOT record trivial inferences. In particular
  "Rs", "Rs.", the rupee sign and "INR" all mean INR - treat them as INR and never raise a currency
  assumption for a standard rupee marker. Set `assumption` to an empty string when nothing material is inferred.
- UNIT INHERITANCE: when a price is listed with NO unit of its own immediately after another line that
  was priced per a pack unit (e.g. "per 100" or "per box"), the bare price most likely shares that pack
  unit. Set quoted_uom to that inherited pack unit as your best reading, put the RAW price in unit_price
  (do NOT pre-divide by the pack size - the system converts), and set `assumption` to name both readings,
  e.g. "unit not stated; assumed the pack unit inherited from the preceding line; the alternative reading
  is per unit". Use plausibility as the tie-breaker: if the per-unit reading would be absurd for a low-value
  item while the per-pack reading is sensible, the pack unit is almost certainly inherited. Never silently
  default such a line to per-unit without flagging.
- DISCOUNT SCOPE: a discount clause written directly after a line clearly applies to THAT line; do not
  flag it. The only uncertainty is whether the discount also reaches BACK to an EARLIER line listed before
  the clause. If it plausibly could, attach the discount to that earlier line too but set `assumption` on
  the EARLIER line ONLY, e.g. "the discount clause follows the later item; assumed it also covers this
  earlier item, but it may apply only to the item it directly follows". Never flag the line the clause directly follows.
- Only infer when there is real supporting context. If there is no basis, do not guess: leave the field
  null and use MISSING/AMBIGUOUS.

COMMERCIAL TERMS & QUALIFICATION
- Fill the structured commercial_terms object with the supplier's stated whole-quote terms. Use the SAME
  fields for every supplier so they compare like-for-like; set a field to null only when the response
  genuinely does not state it (do not invent a value):
  - warranty: the stated warranty as written, e.g. "3 years OEM" or "12 months". Null if unstated.
  - delivery_days: the stated delivery lead time in days as a single number. For a range, use the MAXIMUM
    (e.g. "14-21 days" -> 21). Null if unstated.
  - payment_days: the stated payment term in days, e.g. "net 45" -> 45. Null if unstated.
  - freight_included: true if freight/shipping is included in the prices, false if charged extra. Null if unstated.
  - taxes_included: true if taxes/GST are included in the prices, false if charged extra. Null if unstated.
  - quote_validity_days: how many days the quote/prices stay valid, e.g. "valid 15 days" -> 15. Null if unstated.
  - availability: a SHORT phrase in the supplier's words ONLY when supply is conditional or restricted,
    e.g. "subject to product availability", "subject to stock". Null when supply is not restricted.
- Do NOT editorialise or decide what is a "risk"; just record the facts above. The system decides which
  terms warrant a buyer flag.
- WHOLE-ORDER / BUNDLE DISCOUNT: if the supplier offers a discount on the TOTAL order that applies only
  when the full RFx / entire order is placed with them (distinct from a per-line volume discount),
  capture it as bundle_discount_pct (the percentage) and bundle_discount_condition (the stated condition).
  Do NOT apply it to any line and do NOT treat it as a per-line discount. Leave bundle_discount_pct null
  when there is no such whole-order discount.
- QUALIFICATION: evaluate EVERY supplied questionnaire item on equal footing. Return one entry in
  qualification.items per questionnaire item, each with the question, a result (MET / NOT_MET /
  UNCLEAR) and short evidence in the supplier's own words. Assess ALL items, including delivery.
  Gating rules:
  - Delivery ("deliver within 30 days of PO"): MET only if the supplier's stated lead time is 30 days
    or less. A stated lead time above 30 days, or a range whose MAXIMUM exceeds 30 days, is NOT_MET.
  - ISO 9001: MET only if the certificate is valid AS OF THE CURRENT DATE supplied below. A present-tense
    affirmation of a current certificate in the CURRENT response (e.g. "ISO 9001 current", "current
    certificate available") is MET. A certificate whose expiry date is before the current date is NOT_MET
    even if renewal has been submitted (note "expired, renewal pending"). If the only evidence is a
    PRIOR-YEAR record (last year's certificate), current validity cannot be confirmed, so mark UNCLEAR.
    If validity is otherwise unaddressed, UNCLEAR.
  - Authorized reseller / OEM warranty / enterprise references: MET when the supplier affirms them with
    evidence. Treat "available on request" or an explicit "yes" as MET. UNCLEAR only if not addressed.
- Do NOT set qualification.status yourself beyond your best judgement; the system derives the overall
  status from the per-item results (any NOT_MET fails the supplier).

IMAGES / SCANS
- When an image is attached, read the ORIGINAL image. Use its visual row/column structure to
  bind descriptions to prices. If OCR text conflicts with the image, trust the image.
- Do not guess when the visual evidence itself is insufficient — use AMBIGUOUS.

EVIDENCE & CONFIDENCE
- evidence: concise source-supported wording showing where the value came from.
- confidence: reliability of the extraction from the source, not whether the supplier wins.

Return the complete structured extraction per the schema."""

ANALYST_SYSTEM = """You are an AI procurement analyst. Answer ONLY from the supplied structured RFx
dataset and historical evidence. Be defensible and explicit about uncertainty.

NORMALIZED NUMBERS
- Each quote already carries normalization.normalized_unit_price_inr: a comparable PRE-TAX INR
  per-unit price computed deterministically by the system. Use THIS field for all comparisons and
  award math. Do not re-derive currency/tax/pack conversions yourself.
- comparable=false, or status MISSING/AMBIGUOUS, means the line is NOT eligible for a lowest-cost
  award unless the buyer explicitly asks to investigate it.
- If normalization.freight_extra is true, the landed cost is higher than shown — flag this; never
  call such a total fully landed.

RULES
- Qualification is a hard filter when the buyer asks for qualified suppliers (qualification.status must be PASS).
- Some values may be buyer-resolved: a quote can carry provenance "carried" (last year's price accepted
  for "same as last year"), "buyer-edited" (a value the buyer entered), or "buyer-confirmed" (e.g. an
  accepted substitute), and a qualification item can be buyer-confirmed. Treat resolved values as valid
  for the award, but when you cite one, note it is buyer-resolved rather than a fresh supplier quote.
- Each vendor carries a structured commercial_terms object (warranty, delivery_days, payment_days,
  freight_included, taxes_included, quote_validity_days, availability) and a system-derived
  commercial_risks list. commercial_risks holds only genuine buyer disadvantages (a quote validity shorter
  than the RFx window, or restricted stock/availability). Surface them whenever you recommend or compare
  that vendor and advise the buyer to confirm the risk before relying on it. A supplier may be cheapest yet
  carry a delivery or availability risk a pricier one does not; weigh that in best-value reasoning. Treat
  payment_days (more is better), warranty and freight/tax as neutral comparison terms, not risks.
- Historical evidence is prior-year context only. NEVER treat a historical price as a current quote or
  include it in an award total unless the current response explicitly confirms that number. For
  "same as last year", identify the historical row only when the match is unambiguous, present it as
  historical, and state that supplier confirmation is still required.
- Respect exact configurations. An alternate configuration is non-comparable unless the dataset
  explicitly establishes equivalence.
- ACCEPTED SUBSTITUTES: a comparable line may be a buyer-accepted substitute, marked
  substitute_accepted=true with spec_mismatch_note describing how it differs from the RFx spec (a lower or
  different configuration). Treat it as award-eligible, but keep it DISTINGUISHABLE from an exact-spec quote;
  never silently call it the exact product. When the buyer sets a conditional preference (for example prefer
  the exact product and use a substitute only if it is cheaper by the margin the buyer states), honor it by
  comparing the substitute's normalized price against the cheapest NON-substitute quote on the same line,
  and recommend accordingly. Always state which awarded lines rely on a substitute.
- For any spec or substitute question, use PRECOMPUTED_AWARD.spec_watch: per affected line it gives the
  cheapest_exact_qualified quote (vendor and price) and every substitute with its price, whether the
  supplier is qualified, and how it differs. Compare a substitute ONLY against cheapest_exact_qualified.
  Never invent the exact-spec baseline or omit the cheapest exact quote, and cite the correct line ID from
  spec_watch rather than guessing it.
- For award math: line total = quantity × normalized_unit_price_inr. Show assumptions, exclusions and
  cite line IDs / vendor names. If a question cannot be answered reliably, say exactly what needs review.

PRECOMPUTED AWARD MATH (authoritative)
- You are given PRECOMPUTED_AWARD: deterministic, authoritative totals computed by the system.
  - split_award = cheapest qualified comparable price per line, with total_inr, per-vendor breakdown,
    and lines_unavailable (lines no qualified supplier can comparably serve).
  - single_supplier_awards = per qualified vendor, subtotal_inr and how many lines it covers/misses.
- When the buyer asks for the cheapest award, a total, a per-vendor subtotal, or split-vs-single,
  USE THESE EXACT NUMBERS and ALWAYS state the total in ₹. Never say you lack totals or that the
  backend does not provide a sum — it does, in PRECOMPUTED_AWARD.
- PRECOMPUTED_AWARD also carries resolution_summary: a per-vendor deterministic tally of buyer-resolved
  lines (carried_lines, buyer_edited_lines, accepted_substitute_lines, marked-unavailable, and resolved
  qualification items) with their line IDs. Whenever you state HOW MANY lines are carried, buyer-edited,
  or substituted for a supplier, cite these exact counts. Never tally the lines yourself.
- SPLIT vs SINGLE DIRECTION: PRECOMPUTED_AWARD.comparison states split_total_inr, the
  cheapest_full_basket_single_supplier, the lowest_cost_strategy, and the exact delta between them. When
  you say one strategy is cheaper/lower/higher than the other, READ IT FROM comparison and state the
  direction it gives. Never infer from memory which number is larger; a bundle discount can make a single
  supplier cheaper than the split even though split uses the lowest per-line prices. Only suppliers that
  serve every line have a full-basket single-supplier total; a supplier missing lines is NOT a valid
  full-basket comparison.
- STAY ON THE QUESTION: answer the specific question the buyer asked and do not preempt a later step. If
  the buyer asks only for the split award, give the split; if only for a single-supplier total, give that.
  Only compare split versus single (and state which is cheaper) when the buyer asks to compare, asks which
  is cheapest, or asks for a recommendation. Do not volunteer that a different award strategy would be
  cheaper when it was not asked for. This never suppresses risk warnings, exclusions, or caveats, which you
  always surface.
- WARRANTY PREMIUM: PRECOMPUTED_AWARD.warranty_alternatives gives the cheapest award (split and single)
  restricted to suppliers that meet a warranty tier, and the exact premium in ₹ and % over the overall
  cheapest option. For any "how much more for a longer / 3-year warranty" question, quote these figures;
  do not compute the warranty-restricted award yourself.
- CARRIED COUNT IN CONTEXT: when you describe how much of the SPLIT award rests on carried or edited
  prices, use split_award.per_vendor.carried_lines / buyer_edited_lines / fresh_quote_lines (the count
  among that vendor's split-winning lines). Do not cite a vendor's TOTAL carried count from
  resolution_summary when discussing the split, or "covers N lines" and "M carried" will look
  contradictory. resolution_summary totals are for the full audit trail, not the split subset.
- FREIGHT / TAX INCLUSION: whether a supplier includes freight or tax is a per-vendor fact in its
  commercial_terms (freight_included, taxes_included), also in PRECOMPUTED_AWARD.commercial_flags. State
  it strictly from those flags. Never claim a supplier charges freight or tax extra unless its flag is
  false, and never say "all suppliers except X" without checking every flag. Treat a longer payment term
  as favourable to the buyer, not a risk or a catch.
- Still surface exclusions and uncertainty (lines_unavailable, freight-extra, ambiguous, unqualified),
  but always report the concrete sum alongside the caveats.
- COMMIT WHEN ASKED TO RECOMMEND: when the buyer explicitly asks for a recommendation or a decision, do
  not merely list the options. Name one specific, defensible choice (supplier and total), give the single
  clearest reason, then state the one condition under which the alternative wins (for example, "unless a
  3-year warranty is mandatory, in which case Vertex at +13.7% is the clean firm option").

OUTPUT
- Return a concise buyer-ready answer, an optional table, an optional bar chart (set chart.type='none'
  when a chart adds nothing), warnings, and a recommended next action."""

LINE_ITEM_SYSTEM = (
    "You are a procurement RFx co-pilot. Convert the buyer's request into a single structured RFx line "
    "for an enterprise employee IT hardware refresh. Infer a sensible category (Laptop, Monitor, Dock, "
    "Keyboard, Mouse, Headset, Webcam, Cable, Adapter, etc.), a precise product description, a unit of "
    "measure (usually 'ea'), and a realistic quantity. Use only what "
    "the buyer implies; do not invent unrelated items."
)


# --------------------------------------------------------------------------- #
# Email + inbox workflow (simulated channel; the AI loops stay real)
# --------------------------------------------------------------------------- #
RFX_TITLE = "2026 Employee IT Hardware Refresh"
APPROVED_VENDORS = list(VENDOR_EMAILS.keys())
UPLOAD_DIR = DATA / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

# Each approved vendor's simulated reply. The email body is real content the AI must read;
# for TechSource the quote lives in the body and the attachment is a PRIOR-YEAR sheet.
_TECHSOURCE_BODY = (VENDOR_DIR / "TechSource_India.txt").read_text(encoding="utf-8", errors="ignore")
INBOUND: dict[str, dict] = {
    "Vertex Systems": {
        "body": "Hi Team,\n\nPleased to submit our quote for the refresh. Full line-item pricing, specs and our "
                "questionnaire responses are in the attached workbook.\n\nRegards,\nVertex Bids",
        "attachments": [VENDOR_DIR / "Vertex_Systems.xlsx"],
    },
    "Northstar IT": {
        "body": "Hello,\n\nOur commercial proposal is attached. Please note the volume discount noted on the "
                "proposal.\n\nThanks,\nNorthstar Sales",
        "attachments": [VENDOR_DIR / "Northstar_IT.pdf"],
    },
    "BluePeak Technologies": {
        "body": "Hi,\n\nQuote attached. Prices are in USD; taxes and freight are extra as noted.\n\nBluePeak Sales",
        "attachments": [VENDOR_DIR / "BluePeak_Technologies.docx"],
    },
    "Orion Office Tech": {
        "body": "Hi,\n\nSharing our rate card. Snapped a photo of the printed sheet on my phone, hope it's "
                "legible.\n\nOrion Desk",
        "attachments": [VENDOR_DIR / "Orion_Office_Tech.jpg"],
    },
    "TechSource India": {
        # The quote is the email body only. Last year's sheet is NOT auto-attached: the buyer supplies
        # it as a resolution step ("provide last year's file") to resolve the "same as last year" lines.
        "body": _TECHSOURCE_BODY,
        "attachments": [],
    },
}

# In-memory session for a single presenter: which vendors were invited, and the files collected.
SESSION: dict[str, Any] = {"selected": [], "sent": False, "received": False, "files": {}}


def slug(name: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")


def public_files() -> dict:
    return {v: [{"name": f["name"], "source": f["source"]} for f in files] for v, files in SESSION["files"].items()}


def build_vendor_bundle(vendor: str) -> tuple[str, list[Path]]:
    """Combine all of a vendor's collected files (or the inbound fallback) into text + attachments."""
    text_blocks: list[str] = []
    attach_paths: list[Path] = []
    collected = SESSION["files"].get(vendor)
    if collected:
        items = [(f["name"], Path(f["path"])) for f in collected]
    else:
        inb = INBOUND[vendor]
        items = [(f"{vendor} — email body", None)] + [(p.name, p) for p in inb["attachments"]]
    for name, path in items:
        if path is None:
            text_blocks.append(f"[email body]\n{INBOUND[vendor]['body']}")
        elif path.suffix.lower() in {".pdf"} | IMAGE_EXTS:
            attach_paths.append(path)
            text_blocks.append(f"[attached file: {name} — read from the attachment]")
        else:
            text_blocks.append(f"[file: {name}]\n{extract_local_text(path)[:20000]}")
    return "\n\n".join(text_blocks), attach_paths


def build_parts_multi(system: str, user: str, attach_paths: list[Path]) -> list[dict]:
    parts = [text_part(f"{system}\n\n{user}")]
    for p in attach_paths:
        if p.suffix.lower() in {".pdf"} | IMAGE_EXTS:
            parts.append(file_part(p))
    return parts


def compose_rfx_email(vendor: str) -> dict:
    def _qty(q):
        return int(q) if isinstance(q, (int, float)) and float(q).is_integer() else q
    sample = "\n".join(f"  {it['id']}  {it['description']}  (Qty {_qty(it['quantity'])} {it['uom']})" for it in REQUIREMENTS[:6])
    more = max(0, len(REQUIREMENTS) - 6)
    body = (
        f"Dear {vendor},\n\n"
        f"We invite you to quote for our {RFX_TITLE}.\n\n"
        f"Scope\n{RFX['scope']}\n\n"
        f"Line items ({len(REQUIREMENTS)} total; full spec sheet attached). A sample:\n{sample}\n"
        f"  ...and {more} more.\n\n"
        f"Supplier questionnaire\n" + "\n".join(f"  - {q}" for q in RFX["questionnaire"]) + "\n\n"
        f"Commercial terms\n" + "\n".join(f"  - {t}" for t in RFX["commercial_terms"]) + "\n\n"
        f"Please reply with your quote in whatever format suits you.\n\n"
        f"Regards,\nAcme Procurement"
    )
    return {
        "vendor": vendor,
        "from": "procurement@acme.example",
        "to": VENDOR_EMAILS.get(vendor, ""),
        "subject": f"RFx: {RFX_TITLE}",
        "body": body,
        "attachment": "RFx_line_items.csv",
    }


# --------------------------------------------------------------------------- #
# Request models
# --------------------------------------------------------------------------- #
class CopilotBody(BaseModel):
    message: str = ""


class RfxEdit(BaseModel):
    scope: str | None = None
    questionnaire: list[str] | None = None
    commercial_terms: list[str] | None = None


class SendBody(BaseModel):
    vendors: list[str] | None = None


class AddBody(BaseModel):
    vendors: list[str] | None = None


class RemoveBody(BaseModel):
    vendor: str


class AskBody(BaseModel):
    question: str
    extracted: dict


class ResolveBody(BaseModel):
    extracted: dict
    vendor: str
    line_id: str | None = None
    qual_key: str | None = None
    action: str
    value: Any | None = None
    uom: str | None = None
    note: str | None = None


class ResolveCarriedBody(BaseModel):
    extracted: dict
    vendor: str


class ExportBody(BaseModel):
    extracted: dict


class LineItemBody(BaseModel):
    prompt: str = ""


class LineItemEdit(BaseModel):
    category: str | None = None
    description: str | None = None
    uom: str | None = None
    quantity: float | None = None


# --------------------------------------------------------------------------- #
# Routes
# --------------------------------------------------------------------------- #
@app.get("/api/status")
def status():
    configured = bool(os.getenv("GEMINI_API_KEY")) or (ROOT / "Gemini API Key.txt").exists() or (ROOT.parent / "Gemini API Key.txt").exists()
    return {
        "ai_configured": configured,
        "model": GEMINI_MODEL,
        "provider": "Google Gemini",
        "fx_rate_usd_inr": FX_USD_INR,
        "gst_rate": GST,
    }


@app.get("/api/requirements")
def requirements():
    return {
        "items": REQUIREMENTS,
        "fx_rate_usd_inr": FX_USD_INR,
        "gst_rate": GST,
        "scope": RFX["scope"],
        "questionnaire": RFX["questionnaire"],
        "commercial_terms": RFX["commercial_terms"],
        "historical_vendors": [fy25_vendor()] if fy25_vendor() else [],
    }


@app.put("/api/rfx")
def edit_rfx(body: RfxEdit):
    if body.scope is not None:
        RFX["scope"] = body.scope.strip()
    if body.questionnaire is not None:
        RFX["questionnaire"] = [s.strip() for s in body.questionnaire if s.strip()]
    if body.commercial_terms is not None:
        RFX["commercial_terms"] = [s.strip() for s in body.commercial_terms if s.strip()]
    save_rfx()
    return {"rfx": {k: RFX[k] for k in ("scope", "questionnaire", "commercial_terms")}}


@app.post("/api/rfx-copilot")
def rfx_copilot(body: CopilotBody):
    message = body.message.strip()
    if not message:
        raise HTTPException(400, "Say what you'd like to change in the RFx.")
    state = {
        "scope": RFX["scope"],
        "questionnaire": RFX["questionnaire"],
        "commercial_terms": RFX["commercial_terms"],
        "line_items_count": len(REQUIREMENTS),
    }
    user = f"CURRENT RFx STATE:\n{json.dumps(state, ensure_ascii=False, indent=2)}\n\nBUYER MESSAGE:\n{message}"
    result = gemini_json(build_parts(COPILOT_SYSTEM, user), copilot_schema(), "co-pilot update")

    target = result.get("target")
    added_item = None
    if target == "scope" and result.get("scope"):
        RFX["scope"] = result["scope"].strip()
        save_rfx()
    elif target == "questionnaire" and result.get("questionnaire") is not None:
        RFX["questionnaire"] = [s.strip() for s in result["questionnaire"] if s.strip()]
        save_rfx()
    elif target == "commercial_terms" and result.get("commercial_terms") is not None:
        RFX["commercial_terms"] = [s.strip() for s in result["commercial_terms"] if s.strip()]
        save_rfx()
    elif target == "line_items" and result.get("line_item"):
        li = result["line_item"]
        added_item = {
            "id": next_line_id(),
            "category": str(li.get("category", "")).strip() or "Other",
            "description": str(li.get("description", "")).strip(),
            "uom": str(li.get("uom", "ea")).strip() or "ea",
            "quantity": li.get("quantity") or 0,
        }
        REQUIREMENTS.append(added_item)
        save_requirements()

    return {
        "message": result.get("action_summary") or "Done.",
        "target": target,
        "added_item": added_item,
        "rfx": {k: RFX[k] for k in ("scope", "questionnaire", "commercial_terms")},
        "items": REQUIREMENTS,
    }


@app.post("/api/line-items/generate")
def generate_line_item(body: LineItemBody):
    prompt = body.prompt.strip()
    if not prompt:
        raise HTTPException(400, "Describe the line item to add.")
    drafted = gemini_json(build_parts(LINE_ITEM_SYSTEM, prompt), line_item_schema(), "line item")
    item = {
        "id": next_line_id(),
        "category": str(drafted.get("category", "")).strip() or "Other",
        "description": str(drafted.get("description", "")).strip(),
        "uom": str(drafted.get("uom", "ea")).strip() or "ea",
        "quantity": drafted.get("quantity") or 0,
    }
    REQUIREMENTS.append(item)
    save_requirements()
    return {"item": item, "items": REQUIREMENTS}


@app.put("/api/line-items/{line_id}")
def edit_line_item(line_id: str, body: LineItemEdit):
    for item in REQUIREMENTS:
        if item.get("id") == line_id:
            for field in ("category", "description", "uom", "quantity"):
                value = getattr(body, field)
                if value is not None:
                    item[field] = value
            save_requirements()
            return {"item": item, "items": REQUIREMENTS}
    raise HTTPException(404, f"Line {line_id} not found")


@app.delete("/api/line-items/{line_id}")
def delete_line_item(line_id: str):
    global REQUIREMENTS
    remaining = [item for item in REQUIREMENTS if item.get("id") != line_id]
    if len(remaining) == len(REQUIREMENTS):
        raise HTTPException(404, f"Line {line_id} not found")
    REQUIREMENTS = remaining
    save_requirements()
    return {"items": REQUIREMENTS}


@app.get("/api/approved-vendors")
def approved_vendors():
    return [{"vendor": v, "email": VENDOR_EMAILS.get(v, "")} for v in APPROVED_VENDORS]


def build_inbox() -> list[dict]:
    inbox = []
    for v in SESSION["selected"]:
        inb = INBOUND.get(v)
        if not inb:
            continue
        inbox.append({
            "vendor": v,
            "from": VENDOR_EMAILS.get(v, ""),
            "subject": f"RE: RFx: {RFX_TITLE}",
            "body": inb["body"],
            "attachments": [{"name": p.name, "format": p.suffix[1:].upper()} for p in inb["attachments"]],
            "added": v in SESSION["files"],
        })
    return inbox


@app.get("/api/responses")
def responses_state():
    return {
        "selected": SESSION["selected"],
        "sent": SESSION["sent"],
        "received": SESSION["received"],
        "inbox": build_inbox() if SESSION["received"] else [],
        "files": public_files(),
    }


@app.post("/api/reset-session")
def reset_session():
    SESSION.update({"selected": [], "sent": False, "received": False, "files": {}})
    return {"ok": True}


@app.post("/api/send-rfx")
def send_rfx(body: SendBody):
    """Simulated dispatch. SMTP is stubbed; this composes the buyer-facing RFx email per vendor."""
    targets = [v for v in (body.vendors or APPROVED_VENDORS) if v in APPROVED_VENDORS]
    if not targets:
        raise HTTPException(400, "Select at least one approved vendor.")
    SESSION.update({"selected": targets, "sent": True, "received": False, "files": {}})
    now = datetime.now(timezone.utc)
    emails = {v: compose_rfx_email(v) for v in targets}
    return {
        "sent_at": now.isoformat(),
        "count": len(targets),
        "vendors": [{"vendor": v, "to": VENDOR_EMAILS.get(v, "")} for v in targets],
        "emails": emails,
    }


@app.post("/api/receive-responses")
def receive_responses():
    if not SESSION["selected"]:
        raise HTTPException(400, "Send the RFx first.")
    SESSION["received"] = True
    return {"inbox": build_inbox()}


@app.post("/api/responses/add")
def add_responses(body: AddBody):
    targets = [v for v in (body.vendors or SESSION["selected"]) if v in INBOUND]
    for v in targets:
        inb = INBOUND[v]
        uploads = [f for f in SESSION["files"].get(v, []) if f.get("source") == "uploaded"]  # keep manual uploads
        files = []
        body_path = UPLOAD_DIR / f"{slug(v)}-email-body.txt"
        body_path.write_text(inb["body"], encoding="utf-8")
        files.append({"name": f"{v} email body.txt", "path": str(body_path), "source": "email body"})
        for p in inb["attachments"]:
            files.append({"name": p.name, "path": str(p), "source": "attachment"})
        SESSION["files"][v] = files + uploads
    return {"files": public_files()}


@app.post("/api/responses/upload")
async def upload_response(vendor: str = Form(...), file: UploadFile = File(...)):
    if vendor not in APPROVED_VENDORS:
        raise HTTPException(400, "Unknown vendor.")
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", file.filename or "upload.bin")
    dest = UPLOAD_DIR / f"{slug(vendor)}-{safe}"
    dest.write_bytes(await file.read())
    SESSION["files"].setdefault(vendor, []).append({"name": file.filename or safe, "path": str(dest), "source": "uploaded"})
    return {"files": public_files()}


@app.post("/api/responses/remove")
def remove_response(body: RemoveBody):
    SESSION["files"].pop(body.vendor, None)
    return {"files": public_files()}


@app.get("/api/responses/file")
def responses_file(vendor: str, name: str):
    for f in SESSION["files"].get(vendor, []):
        if f["name"] == name:
            return FileResponse(f["path"], filename=name)
    raise HTTPException(404, "File not found")


@app.post("/api/extract-all")
def extract_all():
    req_text = json.dumps(REQUIREMENTS, indent=2, ensure_ascii=False)
    targets = list(SESSION["files"].keys()) if SESSION["files"] else APPROVED_VENDORS
    vendors = []
    for vendor in targets:
        source, attach_paths = build_vendor_bundle(vendor)
        source = source[:30000]
        user = (
            f"CURRENT DATE: {TODAY} (use this to judge ISO 9001 and certificate validity).\n\n"
            f"RFx requirements (canonical line IDs and requested configurations):\n{req_text}\n\n"
            f"SUPPLIER RFx QUESTIONNAIRE:\n" + "\n".join(f"- {q}" for q in RFX["questionnaire"]) + "\n\n"
            f"FOCUS VENDOR: {vendor}\nSUPPLIER RESPONSE (one or more files/body):\n{source}"
        )
        vendor_data = gemini_json(build_parts_multi(EXTRACT_SYSTEM, user, attach_paths), extraction_schema(), f"extraction for {vendor}")
        vendors.append(enrich_vendor(vendor_data))
    flag_scale_outliers(vendors)
    return {"vendors": vendors, "fx_rate_usd_inr": FX_USD_INR, "gst_rate": GST}


@app.post("/api/resolve-carried")
async def resolve_carried(extracted: str = Form(...), vendor: str = Form(...), note: str = Form(""), file: UploadFile | None = File(None)):
    """Buyer uploads last year's file: resolve 'same as last year' lines into firm carried quotes and
    carry durable qualification answers from THAT file. ISO validity is time-sensitive and stays unresolved."""
    data = json.loads(extracted)
    v = next((x for x in data.get("vendors", []) if x.get("vendor") == vendor), None)
    if not v:
        raise HTTPException(404, "vendor not found")
    if vendor != fy25_vendor():
        raise HTTPException(400, "no prior-year record is held for this supplier")
    rows: list[dict] = []
    if file is not None:
        try:
            rows = _wb_rows(load_workbook(io.BytesIO(await file.read()), data_only=True))
        except Exception:
            rows = []
    prices = prices_from_rows(rows) or fy25_prices()
    quest = questionnaire_from_rows(rows) or fy25_questionnaire()
    src = (file.filename if file and file.filename else (note or "last year's file"))
    resolved = 0
    for q in v.get("quotes", []):
        fy = prices.get(q.get("line_id"))
        if q.get("status") in ("AMBIGUOUS", "MISSING") and fy:
            q["unit_price"] = fy["price"]
            q["quoted_uom"] = fy["uom"]
            q["currency"] = fy["currency"]
            q["tax_included"] = False
            q["freight_included"] = False
            q["status"] = "QUOTED"
            q["provenance"] = "carried"
            q["carried_from"] = src
            q["assumption"] = "carried from last year's price ('same as last year'); confirm with supplier before award"
            q["normalization"] = normalize_quote(q)
            resolved += 1
    # Durable attributes carry; ISO validity is a point-in-time fact and cannot be carried.
    durable = {"reseller": "authorized reseller", "manufacturer": "authorized reseller",
               "warrant": "oem warranty", "referen": "enterprise references"}
    for it in (v.get("qualification") or {}).get("items", []):
        if it.get("result") != "UNCLEAR":
            continue
        ql = (it.get("question") or "").lower()
        if "iso" in ql:
            continue
        ans = next((quest.get(lbl) for kw, lbl in durable.items() if kw in ql), None)
        if ans:
            it["result"] = "MET"
            it["evidence"] = f"{ans} (carried from {src})"
            it["provenance"] = "carried"
    # Carry the durable warranty term into the structured commercial terms when the current
    # response left it blank, so the terms panel is populated consistently for this supplier.
    terms = v.setdefault("commercial_terms", {})
    if not terms.get("warranty") and quest.get("oem warranty"):
        terms["warranty"] = f"{quest['oem warranty']} (carried)"
    derive_commercial_terms(v)
    derive_qual_status(v)
    return {"extracted": data, "resolved_lines": resolved}


@app.post("/api/resolve")
def resolve(body: ResolveBody):
    """Apply one buyer resolution to the dataset and re-run the deterministic engine on the target line."""
    v = next((x for x in body.extracted.get("vendors", []) if x.get("vendor") == body.vendor), None)
    if not v:
        raise HTTPException(404, "vendor not found")
    act = body.action
    if body.line_id:
        q = next((x for x in v.get("quotes", []) if x.get("line_id") == body.line_id), None)
        if not q:
            raise HTTPException(404, "line not found")
        if "original_unit_price" not in q:
            q["original_unit_price"] = q.get("unit_price")
            q["original_uom"] = q.get("quoted_uom")
        if act == "edit_price":
            if body.value is not None:
                q["unit_price"] = float(body.value)
            # A buyer-entered value is the final per-unit INR price; drop the vendor's original
            # unit/currency/discount so the engine does not re-transform it.
            q["quoted_uom"] = body.uom or "ea"
            q["currency"] = "INR"
            q["tax_included"] = False
            q["pack_quantity"] = None
            q["volume_discount_pct"] = None
            q["volume_discount_min_qty"] = None
            q["status"] = "QUOTED"
            q["spec_mismatch"] = False
            q["provenance"] = "buyer-edited"
            q["assumption"] = ""
            q["resolution_note"] = body.note or "value entered by buyer"
        elif act == "set_unit":
            q["quoted_uom"] = body.uom
            q["provenance"] = "buyer-edited"
            q["assumption"] = ""
            q["resolution_note"] = body.note or f"unit set to {body.uom} by buyer"
        elif act == "accept_substitute":
            q["spec_mismatch"] = False
            q["substitute_accepted"] = True  # comparable now, but still an alternate configuration
            q["provenance"] = "buyer-confirmed"
            q["resolution_note"] = body.note or "alternate configuration accepted as equivalent"
        elif act == "reject_substitute":
            q["spec_mismatch"] = True
            q["substitute_accepted"] = False
            q["provenance"] = "buyer-rejected"
            q["resolution_note"] = body.note or "alternate configuration rejected"
        elif act == "approve_assumption":
            q["assumption_confirmed"] = True
            q["provenance"] = "buyer-confirmed"
            q["resolution_note"] = body.note or "assumption approved by buyer"
        elif act == "mark_unavailable":
            q["provenance"] = "buyer-unavailable"
            q["resolution_note"] = body.note or "supplier cannot supply this line"
        elif act == "mark_chase":
            q["resolution_note"] = body.note or "buyer to chase supplier for this line"
        else:
            raise HTTPException(400, f"unknown line action: {act}")
        q["normalization"] = normalize_quote(q)
    elif body.qual_key:
        items = (v.get("qualification") or {}).get("items", [])
        key = body.qual_key.lower()
        target = next((i for i in items if key in (i.get("question") or "").lower()
                       or key in (i.get("evidence") or "").lower()), None)
        if not target:
            raise HTTPException(404, "qualification item not found")
        if act == "confirm_qual":
            target["result"] = body.value or "MET"
            if body.note:
                target["evidence"] = body.note
            target["provenance"] = "buyer-confirmed"
        else:
            raise HTTPException(400, f"unknown qualification action: {act}")
        derive_qual_status(v)
    else:
        raise HTTPException(400, "line_id or qual_key is required")
    return {"extracted": body.extracted}


@app.post("/api/ask")
def ask(body: AskBody):
    dataset = json.dumps(body.extracted, ensure_ascii=False)
    # Historical context now comes from whatever the extractor flagged as prior-year per vendor.
    historical_rows = []
    for v in body.extracted.get("vendors", []):
        for h in (v.get("historical") or []):
            historical_rows.append({"vendor": v.get("vendor"), **h})
    historical = json.dumps(historical_rows, ensure_ascii=False, default=str)
    award = json.dumps(compute_award_context(body.extracted), ensure_ascii=False)
    user = (
        f"CURRENT STRUCTURED RFx DATASET (with system-normalized prices):\n{dataset}\n\n"
        f"PRECOMPUTED_AWARD (authoritative deterministic totals):\n{award}\n\n"
        f"HISTORICAL EVIDENCE (prior-year, not current bids):\n{historical}\n\n"
        f"FX rate INR {FX_USD_INR:.0f}/USD. GST {GST * 100:.0f}%.\n\n"
        f"BUYER QUESTION:\n{body.question}"
    )
    return gemini_json(build_parts(ANALYST_SYSTEM, user), analyst_schema(), "analyst response")


@app.get("/api/source/{vendor}")
def source(vendor: str):
    if vendor not in VENDOR_FILES:
        raise HTTPException(404, "Unknown vendor")
    return FileResponse(VENDOR_FILES[vendor])


@app.get("/api/historical-source/TechSource-India")
def historical_source():
    path = HIST_DIR / "TechSource_India_2025.xlsx"
    if not path.exists():
        raise HTTPException(404)
    return FileResponse(path)


@app.post("/api/export-csv")
def export_csv(body: ExportBody):
    vendors = body.extracted.get("vendors", [])
    by_vendor = {
        v.get("vendor"): {q.get("line_id"): q for q in v.get("quotes", [])}
        for v in vendors
    }
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(
        ["Line", "Description", "Quantity", "UOM"]
        + [f"{v.get('vendor')} (INR/unit, pre-tax)" for v in vendors]
    )
    for item in REQUIREMENTS:
        row = [item["id"], item["description"], item["quantity"], item["uom"]]
        for v in vendors:
            q = by_vendor.get(v.get("vendor"), {}).get(item["id"], {})
            norm = (q.get("normalization") or {}).get("normalized_unit_price_inr")
            row.append("" if norm is None else norm)
        writer.writerow(row)
    data = output.getvalue().encode("utf-8")
    return StreamingResponse(
        io.BytesIO(data),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=kill-the-quote-comparison.csv"},
    )


@app.get("/")
def index():
    return FileResponse(STATIC / "index.html")


@app.get("/{path:path}")
def spa(path: str):
    candidate = STATIC / path
    if candidate.is_file():
        return FileResponse(candidate)
    return FileResponse(STATIC / "index.html")
